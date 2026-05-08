"""Tests for the Flask web layer: rate limit, paywall, unlock token."""

from __future__ import annotations

import io
import os

import docx
import pytest

from crystal.web import create_app


@pytest.fixture
def env(monkeypatch):
    monkeypatch.setenv("SECRET_KEY", "test-secret")
    monkeypatch.delenv("STRIPE_SECRET_KEY", raising=False)
    monkeypatch.delenv("STRIPE_PRICE_ID", raising=False)
    monkeypatch.setenv("UNLOCK_KEY", "letmein")


@pytest.fixture
def app(env):
    return create_app()


@pytest.fixture
def client(app):
    return app.test_client()


def _docx_bytes(*, paragraphs=("Chapter 1", "She walked away.")) -> io.BytesIO:
    buf = io.BytesIO()
    document = docx.Document()
    for i, text in enumerate(paragraphs):
        if i == 0:
            document.add_heading(text, level=1)
        else:
            document.add_paragraph(text)
    document.save(buf)
    buf.seek(0)
    return buf


def _upload(client, **extra):
    return client.post(
        "/analyze",
        data={"manuscript": (_docx_bytes(), "book.docx"), "wph": "9300", **extra},
        content_type="multipart/form-data",
    )


def test_first_analysis_is_free(client):
    r = _upload(client)
    assert r.status_code == 200
    assert b"Total words" in r.data or b"Recording time" in r.data


def test_second_analysis_is_blocked(client):
    _upload(client)  # consume free
    r = _upload(client)
    # Redirects to /paywall
    assert r.status_code in (302, 303)
    assert "/paywall" in r.headers["Location"]


def test_paywall_shows_coming_soon_without_stripe(client):
    _upload(client)
    r = client.get("/paywall")
    assert r.status_code == 200
    assert b"launching soon" in r.data
    assert b"$0.99" in r.data


def test_unlock_token_grants_unlimited(client):
    r = client.get("/unlock/letmein", follow_redirects=False)
    assert r.status_code in (302, 303)

    # Now multiple analyses should succeed.
    for _ in range(3):
        r = _upload(client)
        assert r.status_code == 200


def test_unlock_token_rejects_wrong_value(client):
    r = client.get("/unlock/wrong", follow_redirects=False)
    assert r.status_code == 404


def test_index_shows_credits_remaining(client):
    r = client.get("/")
    assert r.status_code == 200
    assert b"1 report left" in r.data


def test_index_shows_gated_banner_after_use(client):
    _upload(client)
    r = client.get("/")
    assert r.status_code == 200
    assert b"used your free analysis" in r.data


def test_checkout_redirects_to_paywall_when_stripe_not_configured(client):
    r = client.post("/checkout", follow_redirects=False)
    assert r.status_code in (302, 303)
    assert "/paywall" in r.headers["Location"]
